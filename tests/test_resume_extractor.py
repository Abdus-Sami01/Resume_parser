import pytest

from app.schemas.candidate import CandidateProfile
from app.services.extraction.resume_extractor import (
    HeuristicResumeExtractor,
    standardize_profile_skills,
)

SAMPLE_RESUME = """\
Jane Doe
Senior Backend Engineer, jane.doe@example.com, +1 415-555-0100
5+ years of experience building services in Python and Go.
Skills: Python, FastAPI, PostgreSQL, Docker, AWS, Kubernetes
B.S. in Computer Science, Stanford University, 2018
Certifications:
AWS Certified Solutions Architect
"""


def test_heuristic_extractor_pulls_contact_info():
    profile = HeuristicResumeExtractor().extract(SAMPLE_RESUME)
    assert profile.name == "Jane Doe"
    assert profile.email == "jane.doe@example.com"
    assert "415-555-0100" in (profile.phone or "")


def test_heuristic_extractor_standardizes_skills():
    profile = HeuristicResumeExtractor().extract(SAMPLE_RESUME)
    assert "python" in profile.skills
    assert "fastapi" in profile.skills
    assert "aws" in profile.skills


def test_heuristic_extractor_estimates_years_of_experience():
    profile = HeuristicResumeExtractor().extract(SAMPLE_RESUME)
    assert profile.total_years_experience == 5.0


def test_heuristic_extractor_pulls_education():
    profile = HeuristicResumeExtractor().extract(SAMPLE_RESUME)
    assert len(profile.education) == 1

    education = profile.education[0]
    assert education.institution == "Stanford University"
    assert education.graduation_year == 2018
    assert "Computer Science" in education.field_of_study


def test_heuristic_extractor_pulls_certifications():
    profile = HeuristicResumeExtractor().extract(SAMPLE_RESUME)
    assert "AWS Certified Solutions Architect" in profile.certifications


def test_standardize_profile_skills_maps_raw_llm_output_onto_the_taxonomy():
    profile = CandidateProfile(name="Test", skills=["React.js", "ReactJS", "Amazon Web Services"])
    standardized = standardize_profile_skills(profile)

    assert standardized.skills.count("react") == 1
    assert "aws" in standardized.skills


def test_standardize_profile_skills_keeps_unknown_skills_rather_than_dropping_them():
    profile = CandidateProfile(name="Test", skills=["Python", "Obscure Internal Tool"])
    standardized = standardize_profile_skills(profile)

    assert "python" in standardized.skills
    assert "obscure internal tool" in standardized.skills


DATED_RESUME = """\
Jane Doe
jane.doe@example.com
Senior Backend Engineer, Acme Corp, Jan 2019 - Dec 2022
Backend Engineer, Beta Systems, 2015 - 2019
Skills: Python, FastAPI, AWS
"""


def test_dated_roles_are_parsed_individually_rather_than_collapsed():
    profile = HeuristicResumeExtractor().extract(DATED_RESUME)
    assert len(profile.experience) == 2

    roles = {e.role: e for e in profile.experience}
    assert roles["Senior Backend Engineer"].company == "Acme Corp"
    assert roles["Backend Engineer"].company == "Beta Systems"


def test_tenures_sum_instead_of_taking_the_longest():
    """A resume written as date ranges carries no 'N years' phrasing to max() over."""
    profile = HeuristicResumeExtractor().extract(DATED_RESUME)

    assert profile.experience[1].years == 4.0
    assert profile.total_years_experience > 7.5


def test_self_reported_years_still_used_when_no_dates_are_present():
    profile = HeuristicResumeExtractor().extract("Bob\n5+ years of experience in Python.\n")
    assert profile.total_years_experience == 5.0


def test_open_ended_role_runs_to_today():
    from datetime import date

    profile = HeuristicResumeExtractor().extract("Ann\nStaff Engineer, Nova Labs, Mar 2021 - Present\n")
    expected = date.today().year + (date.today().month - 1) / 12 - (2021 + 2 / 12)

    assert profile.experience[0].years == pytest.approx(expected, abs=0.01)


def test_resume_with_no_experience_signal_yields_none():
    assert HeuristicResumeExtractor().extract("Cal\nHobbyist.\n").experience == []


def _build_table_docx() -> bytes:
    """Resumes laid out in tables are common, and `Document.paragraphs` skips table cells."""
    pytest.importorskip("docx")
    from io import BytesIO

    from docx import Document

    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane.doe@example.com")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skills:"
    table.cell(0, 1).text = "Python, FastAPI, AWS, Kubernetes"
    table.cell(1, 0).text = "Experience:"
    table.cell(1, 1).text = "Senior Backend Engineer, Acme Corp, Jan 2019 - Dec 2022"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_table_cells_are_not_dropped():
    from app.services.extraction.document_parser import PlainTextFallbackParser

    text = PlainTextFallbackParser().parse(_build_table_docx(), "table.docx")

    assert "Python" in text
    assert "Acme Corp" in text


def test_table_laid_out_resume_still_yields_a_full_profile():
    from app.services.extraction.document_parser import PlainTextFallbackParser

    text = PlainTextFallbackParser().parse(_build_table_docx(), "table.docx")
    profile = HeuristicResumeExtractor().extract(text)

    assert "python" in profile.skills
    assert "kubernetes" in profile.skills
    assert profile.experience[0].role == "Senior Backend Engineer"
    assert profile.experience[0].company == "Acme Corp"
    assert profile.experience[0].years > 3


def test_role_dates_are_kept_not_just_the_duration():
    """Duration alone cannot distinguish six years ending in 2016 from six ending today."""
    profile = HeuristicResumeExtractor().extract(
        "Ann Lee\nBackend Engineer, Acme Corp, Jan 2010 - Dec 2016\n"
    )

    entry = profile.experience[0]
    assert entry.start_year == pytest.approx(2010.0, abs=0.1)
    assert entry.end_year == pytest.approx(2016.92, abs=0.1)
    assert entry.is_current is False


def test_an_open_ended_role_is_marked_current():
    profile = HeuristicResumeExtractor().extract(
        "Ann Lee\nStaff Engineer, Nova Labs, Mar 2021 - Present\n"
    )

    assert profile.experience[0].is_current is True


BULLETED_RESUME = """\
Ann Lee
ann@x.com

Senior Backend Engineer, Acme Corp, Jan 2020 - Present
- Built Python microservices handling 2M requests/day
- Migrated the billing system to PostgreSQL

Data Engineer, Beta Systems, Jan 2014 - Dec 2018
- Maintained Hadoop clusters
- Wrote Kafka consumers in Scala

Skills: Python, PostgreSQL, Kafka
Education: B.S. in Computer Science, Stanford University, 2013
"""


def test_bullets_under_a_role_are_captured():
    """The title says "Engineer"; the bullets say which language and which database."""
    roles = {e.role: e for e in HeuristicResumeExtractor().extract(BULLETED_RESUME).experience}

    assert roles["Senior Backend Engineer"].achievements == [
        "Built Python microservices handling 2M requests/day",
        "Migrated the billing system to PostgreSQL",
    ]


def test_bullets_are_attached_to_the_right_role():
    roles = {e.role: e for e in HeuristicResumeExtractor().extract(BULLETED_RESUME).experience}

    assert any("Kafka" in a for a in roles["Data Engineer"].achievements)
    assert not any("Kafka" in a for a in roles["Senior Backend Engineer"].achievements)


def test_a_following_section_does_not_get_swallowed_as_a_bullet():
    """"Skills:" and "Education:" carry content, so a bare-header check misses them."""
    profile = HeuristicResumeExtractor().extract(BULLETED_RESUME)

    for entry in profile.experience:
        assert not any(a.startswith("Skills") or a.startswith("Education") for a in entry.achievements)
    assert profile.education  # the education line was still parsed as education


def test_wrapped_bullets_are_joined():
    profile = HeuristicResumeExtractor().extract(
        "Ann\nEngineer, Acme, Jan 2020 - Present\n"
        "- Built a service that processed\n"
        "  several million events per day\n"
    )

    assert profile.experience[0].achievements == [
        "Built a service that processed several million events per day"
    ]


def test_a_role_without_bullets_yields_no_achievements():
    profile = HeuristicResumeExtractor().extract("Ann\nEngineer, Acme, Jan 2020 - Present\n")
    assert profile.experience[0].achievements == []
